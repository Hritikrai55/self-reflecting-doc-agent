"""
Word Document Generator — creates professionally styled .docx files.

Uses python-docx to produce documents with:
- Title page
- Consistent typography (Calibri)
- Styled headings, body text, bullet lists, and tables
- Page numbers and metadata
"""

from __future__ import annotations

import os
import re
from datetime import datetime

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ── Colour palette ───────────────────────────────────────────────────

PRIMARY = RGBColor(0x1A, 0x3C, 0x6E)       # Deep navy
SECONDARY = RGBColor(0x3A, 0x7C, 0xBD)     # Medium blue
ACCENT = RGBColor(0x2E, 0x86, 0xAB)        # Teal accent
TEXT_DARK = RGBColor(0x2D, 0x2D, 0x2D)      # Near black
TEXT_LIGHT = RGBColor(0x5A, 0x5A, 0x5A)     # Grey
TABLE_HEADER_BG = "1A3C6E"                  # hex for table shading


def generate_docx(
    title: str,
    sections: list[dict[str, str]],
    document_type: str,
    output_dir: str = "outputs",
) -> str:
    """
    Generate a polished .docx file from the assembled sections.

    Parameters
    ----------
    title : str
        Document title.
    sections : list[dict]
        Each dict has keys 'heading' and 'content'.
    document_type : str
        The type of document (shown on title page).
    output_dir : str
        Directory to save the file in.

    Returns
    -------
    str
        The filename of the generated document.
    """
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()

    # ── Page setup ───────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # ── Default font ─────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = TEXT_DARK

    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.15

    # ── Heading styles ───────────────────────────────────────────
    for level, (size, color, bold) in enumerate(
        [
            (Pt(26), PRIMARY, True),    # Heading 1
            (Pt(18), SECONDARY, True),  # Heading 2
            (Pt(14), ACCENT, True),     # Heading 3
        ],
        start=1,
    ):
        h_style = doc.styles[f"Heading {level}"]
        h_font = h_style.font
        h_font.name = "Calibri"
        h_font.size = size
        h_font.color.rgb = color
        h_font.bold = bold

    # ── Title page ───────────────────────────────────────────────
    _add_title_page(doc, title, document_type)

    # ── Content sections ─────────────────────────────────────────
    for sec in sections:
        _add_section(doc, sec["heading"], sec["content"])

    # ── Footer with page numbers ─────────────────────────────────
    _add_page_numbers(doc)

    # ── Save ─────────────────────────────────────────────────────
    safe_title = re.sub(r"[^\w\s-]", "", title)[:50].strip().replace(" ", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{safe_title}_{timestamp}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    return filename


# ── Internal helpers ─────────────────────────────────────────────────

def _add_title_page(doc: Document, title: str, document_type: str) -> None:
    """Add a styled title page."""
    # Spacer
    for _ in range(4):
        doc.add_paragraph()

    # Document type label
    p_type = doc.add_paragraph()
    p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_type.add_run(document_type.upper())
    run.font.size = Pt(14)
    run.font.color.rgb = ACCENT
    run.font.bold = True
    run.font.name = "Calibri"

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title)
    run.font.size = Pt(32)
    run.font.color.rgb = PRIMARY
    run.font.bold = True
    run.font.name = "Calibri"

    # Divider line
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_div.add_run("━" * 40)
    run.font.color.rgb = SECONDARY
    run.font.size = Pt(12)

    # Date
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_date.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(12)
    run.font.color.rgb = TEXT_LIGHT
    run.font.name = "Calibri"

    # Attribution
    p_attr = doc.add_paragraph()
    p_attr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_attr.add_run("Prepared by Autonomous AI Agent")
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT_LIGHT
    run.font.italic = True
    run.font.name = "Calibri"

    # Page break after title
    doc.add_page_break()


def _add_section(doc: Document, heading: str, content: str) -> None:
    """Parse content and add it with appropriate formatting."""
    # Add section heading
    doc.add_heading(heading, level=1)

    # Process content line by line
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Table block
        if line == "TABLE_START":
            table_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != "TABLE_END":
                table_lines.append(lines[i].strip())
                i += 1
            i += 1  # skip TABLE_END
            if table_lines:
                _add_table(doc, table_lines)
            continue

        # Markdown-style table (lines with | separators)
        if "|" in line and i + 1 < len(lines) and "|" in lines[i + 1]:
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                stripped = lines[i].strip()
                # Skip separator lines (e.g., |---|---|)
                if not re.match(r"^[\|\s\-:]+$", stripped):
                    table_lines.append(stripped)
                i += 1
            if table_lines:
                _add_table(doc, table_lines)
            continue

        # Sub-heading (## or ### or **Heading**)
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
            i += 1
            continue

        # Bullet point
        if line.startswith(("- ", "• ", "* ")):
            text = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_run(p, text)
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+[\.\)]\s", line):
            text = re.sub(r"^\d+[\.\)]\s*", "", line)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_run(p, text)
            i += 1
            continue

        # Bold heading line (e.g., "**Something Important**")
        bold_match = re.match(r"^\*\*(.+?)\*\*$", line)
        if bold_match:
            doc.add_heading(bold_match.group(1), level=2)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_formatted_run(p, line)
        i += 1


def _add_formatted_run(paragraph, text: str) -> None:
    """Add text to a paragraph, handling inline **bold** and *italic* markdown."""
    # Split on bold markers
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = "Calibri"
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = "Calibri"
        else:
            run = paragraph.add_run(part)
            run.font.name = "Calibri"


def _add_table(doc: Document, rows: list[str]) -> None:
    """
    Add a styled table from pipe-separated rows.
    First row is treated as the header.
    """
    if not rows:
        return

    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)

    if not parsed:
        return

    num_cols = max(len(r) for r in parsed)

    # Pad rows with fewer columns
    for r in parsed:
        while len(r) < num_cols:
            r.append("")

    table = doc.add_table(rows=len(parsed), cols=num_cols, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, row_data in enumerate(parsed):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.name = "Calibri"
            run.font.size = Pt(10)

            if row_idx == 0:
                # Header row styling
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shading = cell._element.get_or_add_tcPr()
                shading_elm = shading.makeelement(
                    qn("w:shd"),
                    {
                        qn("w:fill"): TABLE_HEADER_BG,
                        qn("w:val"): "clear",
                    },
                )
                shading.append(shading_elm)

    # Add spacing after table
    doc.add_paragraph()


def _add_page_numbers(doc: Document) -> None:
    """Add page numbers to the document footer."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Page number field
    run = p.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_LIGHT
    run.font.name = "Calibri"

    fldChar1 = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    run._element.append(fldChar1)

    instrText = run._element.makeelement(qn("w:instrText"), {})
    instrText.text = " PAGE "
    run._element.append(instrText)

    fldChar2 = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run._element.append(fldChar2)
